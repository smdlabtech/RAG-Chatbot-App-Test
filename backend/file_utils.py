import os
import hashlib
import tempfile
from docx import Document as DocxDocument
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

def get_file_hash(file_bytes):
    return hashlib.md5(file_bytes).hexdigest()

def get_title_from_filename(filename):
    return os.path.splitext(os.path.basename(filename))[0]

def generate_export_file(data, format="txt"):
    answer = data.get("answer", "Aucune réponse")
    context = data.get("context", [])

    if format == "txt":
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8")
        temp.write("Réponse générée :\n")
        temp.write(answer + "\n\n")
        if context:
            temp.write("Contexte utilisé :\n")
            for i, doc in enumerate(context):
                temp.write(f"Chunk {i+1} : {doc.page_content}\n---\n")
        temp.close()
        return temp.name

    elif format == "docx":
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc = DocxDocument()
        doc.add_heading("Réponse générée", 0)
        doc.add_paragraph(answer)
        if context:
            doc.add_heading("Contexte utilisé", level=1)
            for i, doc_chunk in enumerate(context):
                doc.add_paragraph(f"Chunk {i+1} :", style='Heading2')
                doc.add_paragraph(doc_chunk.page_content)
        doc.save(temp.name)
        return temp.name

    elif format == "pdf":
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        doc = SimpleDocTemplate(temp.name)
        styles = getSampleStyleSheet()
        flowables = [
            Paragraph("Réponse générée", styles["Heading1"]),
            Paragraph(answer, styles["Normal"]),
        ]
        if context:
            flowables.append(Paragraph("Contexte utilisé", styles["Heading2"]))
            for i, doc_chunk in enumerate(context):
                flowables.append(Paragraph(f"Chunk {i+1} :", styles["Heading3"]))
                flowables.append(Paragraph(doc_chunk.page_content, styles["Normal"]))
        doc.build(flowables)
        return temp.name

    else:
        raise ValueError("Format non pris en charge")
